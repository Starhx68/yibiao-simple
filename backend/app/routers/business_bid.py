import os
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File
from typing import List, Optional
from pydantic import BaseModel
from ..services.auth_service import get_current_user
from ..models.models import (
    User, BusinessBidProject, BusinessBidDirectory, BusinessBidElement,
    CompanyInfo, Qualification, Personnel, FinancialInfo, Performance
)
from ..database import get_db
from sqlalchemy.orm import Session
from ..services.file_service import FileService
from ..services.openai_service import OpenAIService
from ..utils.config_manager import config_manager
from ..utils.sse import sse_response
import uuid
import json
import logging
import re
import time
from html import escape
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from app.config import settings

# 配置日志
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

router = APIRouter(prefix="/api/business-bids", tags=["商务标管理"])

class BusinessProjectCreate(BaseModel):
    project_name: str

@router.post("/")
def create_business_project(data: BusinessProjectCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project_id = str(uuid.uuid4())
    project = BusinessBidProject(
        id=project_id,
        user_id=current_user.id,
        project_name=data.project_name,
        status="draft"
    )
    db.add(project)
    db.commit()
    db.refresh(project)
    return {"id": project.id, "project_name": project.project_name, "status": project.status}

@router.get("/")
def list_business_projects(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    projects = db.query(BusinessBidProject).filter(BusinessBidProject.user_id == current_user.id).all()
    return {"items": [{"id": p.id, "project_name": p.project_name, "status": p.status, "created_at": p.created_at} for p in projects]}

@router.post("/{project_id}/upload-tender")
async def upload_tender_document(project_id: str, file: UploadFile = File(...), current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    try:
        # 复用 FileService 读取文件内容，并获取保存路径
        file_content, file_path = await FileService.process_uploaded_file(file)
        
        # 计算相对路径，用于前端访问
        # 假设 settings.upload_dir 是绝对路径或相对于项目根目录的路径
        # 我们需要将其转换为相对于 static mount point 的路径
        # app.mount("/uploads", StaticFiles(directory=settings.upload_dir), name="uploads")
        # file_path e.g. "D:/Temp/bsdoc-sw/hxybs-js/backend/uploads/filename_timestamp.docx"
        # relative_path should be "filename_timestamp.docx"
        
        filename = os.path.basename(file_path)
        # 构造可访问的URL (假设后端服务运行在同域或已配置代理)
        # 前端可以通过 /uploads/{filename} 访问
        tender_document_url = f"/uploads/{filename}"

        # 保存到数据库
        project.tender_document_name = file.filename
        project.tender_document_url = tender_document_url
        project.tender_content = file_content
        project.status = "analyzing"
        db.commit()
        return {"success": True, "message": "File uploaded and parsed successfully", "file_content": file_content[:1000] + "..."}
    except Exception as e:
        logger.error(f"File upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{project_id}")
def get_business_project(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    # 构建文件的可访问URL
    tender_document_url = None
    if project.tender_document_name:
        # 假设文件保存在 settings.upload_dir 目录下，并且文件名中包含时间戳
        # 我们需要找到实际的文件名。但在 upload_tender_document 中我们只存了原始文件名到 tender_document_name
        # 实际上 FileService.save_uploaded_file 生成了新的文件名。
        # 这里为了简化，我们暂时只返回原始文件名，前端可能需要适配。
        # 更好的做法是在数据库中存储实际保存的文件路径或文件名。
        # 暂时 workaround: 我们需要让 upload 接口把实际路径存下来。
        pass

    return {
        "id": project.id,
        "project_name": project.project_name,
        "status": project.status,
        "tender_document_name": project.tender_document_name,
        "tender_document_url": project.tender_document_url, # 新增字段
        "tender_content": project.tender_content,
        "created_at": project.created_at
    }

@router.post("/{project_id}/analyze-stream")
async def analyze_business_bid_stream(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    if not project.tender_content:
        raise HTTPException(status_code=400, detail="Tender document not uploaded or content is empty")

    config = config_manager.load_config()
    if not config.get('api_key'):
        raise HTTPException(status_code=400, detail="请先配置OpenAI API密钥")

    openai_service = OpenAIService()

    async def generate():
        from app.database import SessionLocal
        import asyncio
        system_prompt = """你是一名专业的招标文件商务标分析师，擅长从复杂的招标文档中提取关键信息。
请你严格按照以下一级分类、二级分类，从提供的招标文件中，精准提取对应二级分类下的全部关键信息，要求信息完整、表述精准、层次分明，若招标文件中某分类下无相关信息，需明确标注“无相关要求”：

# 一、投标人资格要求
- 主体资格要求：投标人主体类型、主体证明文件要求、营业执照相关要求、外资/合资企业相关备案证明要求、关联企业投标限制条件。
- 资质与许可要求：核心资质要求、安全生产相关许可要求、行业专项资质要求、所有资质证书的有效性要求。
- 财务状况要求：财务报表要求、财务指标要求、银行资信证明/资金证明要求、纳税证明要求、社保缴纳证明要求、履约能力相关限制。
- 信誉要求：失信记录要求、近3-5年重大违法记录要求、履约评价要求、法定代表人及项目负责人行贿犯罪记录要求。
- 业绩要求：类似项目要求、业绩数量要求、业绩质量要求、业绩证明材料要求、业绩限制条件。
- 人员配置要求：项目负责人要求、技术负责人要求、核心团队成员要求、人员稳定性要求。
- 其他资格要求：联合体投标要求、代理商投标要求、本地化服务要求、合规投标承诺要求。

# 二、投标费用与投标保证金
- 投标费用：投标相关费用承担方、招标文件售价、缴纳方式、售后是否可退、踏勘现场等其他费用承担方。
- 投标保证金：保证金具体金额、缴纳方式、指定缴纳账户、缴纳截止时间、保证金退还规则、保证金不予退还情形。

# 三、投标文件商务部分编制要求
- 编制原则：响应性要求、真实性要求、完整性要求。
- 编制内容与格式：商务部分组成内容及排列顺序、纸张规格、装订要求、签署盖章要求、正副本份数、电子版要求。
- 报价相关编制要求：报价货币、报价范围、报价形式、最高/最低限价要求、投标有效期、报价是否可修改。

# 四、开标与评标中的商务评审要点
- 开标阶段商务核查：投标文件密封核查要求、资格初步核查内容、报价核查内容。
- 评标阶段商务评审：资格详细评审规则、商务评分项具体内容及分值设置、商务偏离评审规则。

# 五、中标与合同相关商务条款
- 中标相关商务要求：中标公示期限、异议处理方式、中标通知书发放、履约保证金要求、中标人义务。
- 合同主要商务条款：合同范围与内容、合同价款形式及金额、付款方式、质保金、价格调整、交付要求、验收标准、质保期、违约责任、争议解决。

# 六、商务部分其他关键要求
- 保密要求：保密范围、保密责任。
- 投标文件的修改与撤回：修改规则、撤回规则。
- 废标相关商务情形：所有商务类废标情形。
- 特殊项目商务补充要求：工程类/服务类/货物类的专项商务要求。

# 七、关键信息与占位符映射
- 提取用于投标文件的关键实体信息，并设定统一的占位符。例如：项目名称对应 {{PROJECT_NAME}}，招标人对应 {{TENDERER_NAME}}，投标人名称对应 {{COMPANY_NAME}}，法定代表人对应 {{LEGAL_REP_NAME}}，投标总价对应 {{BID_PRICE}}，投标有效期对应 {{BID_VALIDITY}} 等。

必须严格输出为JSON数组格式，结构如下：
[
  {
    "title": "投标人资格要求",
    "subcategories": [
      {
        "title": "主体资格要求",
        "items": [
          {"name": "投标人主体类型", "description": "..."},
          {"name": "营业执照相关要求", "description": "..."}
        ]
      }
    ]
  },
  {
    "title": "关键信息与占位符映射",
    "subcategories": [
      {
        "title": "项目与企业信息",
        "items": [
          {"name": "项目名称", "placeholder": "{{PROJECT_NAME}}", "value": "提取到的实际项目名称，例如：XXX采购项目"},
          {"name": "招标人", "placeholder": "{{TENDERER_NAME}}", "value": "提取到的实际招标人名称"}
        ]
      }
    ]
  }
]

必须直接输出JSON字符串，不能包含 ```json 等任何Markdown代码块标记，也不要包含任何其他说明文字。
"""
        # 使用切片后的文本，但为了避免截断关键信息，建议尽量使用全文。
        # 如果文本过长，可能需要分段处理或使用支持长上下文的模型。
        # 这里假设 150k 字符对于大多数标书足够，且模型支持。
        user_prompt = f"请分析以下招标文件内容，提取商务相关要素：\n\n{project.tender_content[:150000]}"
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        db_gen = SessionLocal()
        try:
            full_response = ""
            try:
                async for chunk in openai_service.stream_chat_completion(messages, temperature=0.1):
                    full_response += chunk
                    yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
            except asyncio.CancelledError:
                logger.warning("Client disconnected during streaming. Saving partial result.")
                # We will still try to save the partial result below
            except Exception as inner_e:
                logger.error(f"Error during OpenAI stream: {inner_e}")
                full_response += f"\n[Error: {str(inner_e)}]"
                yield f"data: {json.dumps({'error': str(inner_e)}, ensure_ascii=False)}\n\n"

            logger.info(f"Analysis result for project {project_id}: {full_response}")

            # 尝试修复可能的JSON格式错误（例如包含Markdown标记）
            cleaned_response = full_response.strip()
            if cleaned_response.startswith("```json"):
                cleaned_response = cleaned_response[7:]
            if cleaned_response.endswith("```"):
                cleaned_response = cleaned_response[:-3]
            cleaned_response = cleaned_response.strip()

            project_gen = db_gen.query(BusinessBidProject).filter(BusinessBidProject.id == project_id).first()
            if project_gen:
                # 验证JSON合法性
                try:
                    json.loads(cleaned_response)
                    # Save the result to DB
                    project_gen.elements_content = cleaned_response
                    project_gen.status = "analyzed"
                    db_gen.commit()
                    logger.info("Successfully saved analyzed elements to DB.")
                except json.JSONDecodeError as e:
                    logger.error(f"JSON decode error: {e}, response: {cleaned_response}")
                    # 即使解析失败，也保存原始响应以便排查，或者标记为失败
                    project_gen.elements_content = cleaned_response 
                    project_gen.status = "analyzed_with_error"
                    db_gen.commit()
                    logger.info("Saved analyzed_with_error to DB.")
            
            yield "data: [DONE]\n\n"
        except asyncio.CancelledError:
            logger.warning("Client disconnected. Analysis aborted in outer block.")
        except Exception as e:
            logger.error(f"Analysis failed: {e}")
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"
            yield "data: [DONE]\n\n"
        finally:
            db_gen.close()

    return sse_response(generate())

@router.get("/{project_id}/elements")
def get_business_elements(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    default_elements = []
    
    if project.elements_content:
        try:
            elements = json.loads(project.elements_content)
            # 确保每个项都有id
            if isinstance(elements, list):
                for cat_idx, category in enumerate(elements):
                    category["id"] = f"cat_{cat_idx}"
                    if "subcategories" in category:
                        for sub_idx, subcat in enumerate(category["subcategories"]):
                            subcat["id"] = f"sub_{cat_idx}_{sub_idx}"
                            if "items" in subcat:
                                for item_idx, item in enumerate(subcat["items"]):
                                    item["id"] = f"item_{cat_idx}_{sub_idx}_{item_idx}"
            return elements
        except json.JSONDecodeError:
            pass
            
    return default_elements

from app.database import SessionLocal

def find_title_pos(text: str, search_title: str, start_pos: int = 0, end_pos: int = -1) -> int:
    if not search_title:
        return -1
    if end_pos == -1:
        end_pos = len(text)
        
    search_text = text[start_pos:end_pos]
    
    # 去除标题中可能包含的编号（如 "1. ", "1.1 ", "一、" 等）提取核心部分
    core_title = re.sub(r"^[一二三四五六七八九十\d\.\s、]+", "", search_title).strip()
    if not core_title:
        core_title = search_title
        
    # 尝试精确匹配带有核心标题的行
    pattern = re.compile(rf"(?m)^[ \t]*[一二三四五六七八九十\d\.\s、]*{re.escape(core_title)}[ \t]*$")
    m = pattern.search(search_text)
    if m:
        return start_pos + m.start()
        
    # 降级：部分匹配
    idx = search_text.find(core_title)
    if idx != -1:
        # 尝试回退到该行的行首
        last_newline = search_text.rfind('\n', 0, idx)
        return start_pos + (last_newline + 1 if last_newline != -1 else idx)
        
    return -1

def get_format_chapter_info(content_text: str):
    """
    提取“投标文件格式”或“招标文件格式”的章节文本和位置
    返回: (format_chapter_pos, format_chapter_end, chapter_text)
    如果没有找到，返回 (-1, -1, None)
    """
    if not content_text:
        return -1, -1, None
        
    format_chapter_pos = -1
    format_chapter_end = len(content_text)
    
    format_match = re.search(r"(?m)^[ \t]*(?:第[一二三四五六七八九十百]+[章部分][ \t]+)?.*(?:招标|投标)文件.*格式.*[ \t]*$", content_text)
    if format_match:
        format_chapter_pos = format_match.start()
    else:
        fallback_idx = content_text.find("投标文件格式")
        if fallback_idx == -1:
            fallback_idx = content_text.find("招标文件格式")
        if fallback_idx != -1:
            format_chapter_pos = fallback_idx

    if format_chapter_pos >= 0:
        next_chapter_match = re.search(r"(?m)^[ \t]*第[一二三四五六七八九十百]+[章部分][ \t]+", content_text[format_chapter_pos + 10:])
        if next_chapter_match:
            format_chapter_end = format_chapter_pos + 10 + next_chapter_match.start()
        return format_chapter_pos, format_chapter_end, content_text[format_chapter_pos:format_chapter_end]
        
    return -1, -1, None

def generate_format_word_file(project, db: Session):
    """
    根据项目的招标文档，截取“投标文件格式”章节，并生成新的Word文件。
    将新生成的Word文件URL保存到 project.other_urls (以JSON形式)。
    """
    if not project.tender_content:
        return

    content_text = project.tender_content
    pos, end, chapter_text = get_format_chapter_info(content_text)
    
    if not chapter_text:
        return
    
    # 确定输出路径
    import time
    filename = f"format_{int(time.time())}.docx"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(settings.upload_dir, exist_ok=True)
    
    # 如果原文件是 docx，尝试直接裁剪原文件以保留真实格式
    original_url = project.tender_document_url
    if original_url and original_url.endswith('.docx'):
        original_filename = original_url.split('/')[-1]
        original_path = os.path.join(settings.upload_dir, original_filename)
        if os.path.exists(original_path):
            try:
                import docx
                doc = docx.Document(original_path)
                in_chapter = False
                format_chapter_found = False
                
                # 遍历所有 block-level 元素
                for element in list(doc.element.body):
                    if element.tag.endswith('sectPr'):
                        continue
                        
                    if element.tag.endswith('p'):
                        from docx.text.paragraph import Paragraph
                        p = Paragraph(element, doc)
                        text = p.text.strip()
                        
                        if not in_chapter:
                            if re.match(r"^(?:第[一二三四五六七八九十百]+[章部分][ \t]+)?.*(?:招标|投标)文件.*格式.*$", text):
                                in_chapter = True
                                format_chapter_found = True
                            else:
                                element.getparent().remove(element)
                        else:
                            if re.match(r"^第[一二三四五六七八九十百]+[章部分][ \t]+", text):
                                in_chapter = False
                                element.getparent().remove(element)
                    else:
                        if not in_chapter:
                            element.getparent().remove(element)
                
                if format_chapter_found:
                    doc.save(output_path)
                    new_url = f"/uploads/{filename}"
                    try:
                        other_urls = json.loads(project.other_urls) if project.other_urls else {}
                    except:
                        other_urls = {}
                    other_urls["format_document_url"] = new_url
                    project.other_urls = json.dumps(other_urls, ensure_ascii=False)
                    db.commit()
                    return
            except Exception as e:
                logger.error(f"Failed to crop original docx: {e}")

    # 降级：从提取的纯文本创建一个新的 Word 文档
    doc = docx.Document()
    try:
        styles = doc.styles
        base_styles = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]
        for style_name in base_styles:
            if style_name in styles:
                style = styles[style_name]
                font = style.font
                font.name = "宋体"
                if style._element.rPr is None:
                    style._element._add_rPr()
                rpr = style._element.rPr
                rpr.rFonts.set(qn("w:eastAsia"), "宋体")
                if style_name == "Normal":
                    font.bold = False
    except Exception:
        pass

    # 简单的 Markdown 解析插入段落
    lines = chapter_text.split('\n')
    for line in lines:
        if not line.strip():
            continue
        p = doc.add_paragraph()
        run = p.add_run(line.strip())
        run.font.name = "宋体"
        r = run._element.rPr
        if r is not None and r.rFonts is not None:
            r.rFonts.set(qn("w:eastAsia"), "宋体")
            
    # 保存新的 Word 文档
    doc.save(output_path)
    
    # 更新 project 的 other_urls
    new_url = f"/uploads/{filename}"
    
    try:
        other_urls = json.loads(project.other_urls) if project.other_urls else {}
    except:
        other_urls = {}
        
    other_urls["format_document_url"] = new_url
    project.other_urls = json.dumps(other_urls, ensure_ascii=False)
    db.commit()


def extract_content_for_directories(directories, project):
    from bs4 import BeautifulSoup
    import mammoth
    import os
    import json
    import re
    from app.config import settings

    content_text = project.tender_content or ""
    format_chapter_pos, format_chapter_end, chapter_text = get_format_chapter_info(content_text)
    
    # 尝试加载 HTML 格式的格式章节
    html_elements = []
    use_html = False
    
    # 优先尝试从原文件提取 HTML，避免裁切后的 docx 损坏导致 mammoth 崩溃
    docx_url = project.tender_document_url
    if docx_url and docx_url.endswith('.docx'):
        filename = docx_url.split('/')[-1]
        docx_path = os.path.join(settings.upload_dir, filename)
        if os.path.exists(docx_path):
            try:
                with open(docx_path, "rb") as f:
                    style_map = """
                    p[style-name='Heading 1'] => h1:fresh
                    p[style-name='Heading 2'] => h2:fresh
                    p[style-name='Heading 3'] => h3:fresh
                    table => table:fresh
                    u => u
                    """
                    result = mammoth.convert_to_html(f, style_map=style_map)
                    html = result.value
                soup = BeautifulSoup(html, 'html.parser')
                html_elements = soup.find_all(recursive=False)
                if html_elements:
                    use_html = True
            except Exception as e:
                logger.error(f"Failed to parse original docx to HTML with mammoth: {e}")

    flat_nodes = []
    def flatten(nodes):
        for n in nodes:
            flat_nodes.append(n)
            if n.get("children"):
                flatten(n["children"])
    flatten(directories)
    
    def clean_title(t):
        t = re.sub(r"^[一二三四五六七八九十\d\.\s、]+", "", t).strip()
        return t

    # 找到 html 中的格式章节起始位置，跳过目录中的干扰项
    html_format_chapter_idx = 0
    html_format_chapter_end_idx = 0
    
    if use_html:
        html_format_chapter_end_idx = len(html_elements)
        for idx, el in enumerate(html_elements):
            text = el.get_text().strip()
            # 必须是独立的标题行，不能太长
            if re.match(r"^(?:第[一二三四五六七八九十百]+[章部分][ \t]+)?.*(?:招标|投标)文件.*格式.*$", text) and len(text) < 50:
                # 为了防止匹配到目录中的“第八章 投标文件格式 ........... 150”或包含制表符加页码的
                if not re.search(r"(\.{3,}|\t)\s*\d+", text):
                    html_format_chapter_idx = idx
                    # 不 break，取最后一个匹配的（通常真实章节在目录之后）
        
        # 寻找格式章节的结束位置（即下一个大章节的起始位置）
        if html_format_chapter_idx > 0:
            for idx in range(html_format_chapter_idx + 1, len(html_elements)):
                text = html_elements[idx].get_text().strip()
                # 匹配 第X章 或 第X部分
                if re.match(r"^第[一二三四五六七八九十百]+[章部分][ \t]+", text) and len(text) < 50:
                    if not re.search(r"(\.{3,}|\t)\s*\d+", text):
                        html_format_chapter_end_idx = idx
                        break

    for i, node in enumerate(flat_nodes):
        title = node.get("title", "")
        next_title = None
        if i + 1 < len(flat_nodes):
            next_title = flat_nodes[i + 1].get("title", "")
            
        core_title = clean_title(title)
        
        if use_html:
            core_next_title = clean_title(next_title) if next_title else None
            
            start_idx = -1
            # 放宽匹配条件，只要包含即可，不再限制长度，因为有些标题会被包裹在很长的格式段落中
            # 仅在“投标文件格式”章节内查找
            for idx in range(html_format_chapter_idx, html_format_chapter_end_idx):
                el = html_elements[idx]
                text = el.get_text().strip()
                # 针对标题行，通常不会太长，但为了容错，把长度限制放宽到 100
                if core_title and core_title in text and len(text) < len(core_title) + 100:
                    start_idx = idx
                    break
            
            if start_idx != -1:
                end_idx = html_format_chapter_end_idx
                if core_next_title:
                    for idx in range(start_idx + 1, html_format_chapter_end_idx):
                        text = html_elements[idx].get_text().strip()
                        if core_next_title in text and len(text) < len(core_next_title) + 100:
                            end_idx = idx
                            break
                
                # 拼接 HTML
                html_parts = []
                for idx in range(start_idx, end_idx):
                    # 将 html 字符串转存
                    el_str = str(html_elements[idx])
                    html_parts.append(el_str)
                    
                # 过滤掉一些可能导致编辑器报错的无意义标签，比如空锚点
                raw_html = ''.join(html_parts)
                raw_html = re.sub(r'<a id="[^"]+"></a>', '', raw_html)
                
                # 为保证原样显示，补充必要的tbody标签，并增加表格边框，防止WangEditor过滤或无边框
                raw_html = raw_html.replace('<table>', '<table border="1" width="100%"><tbody>')
                raw_html = raw_html.replace('</table>', '</tbody></table>')
                
                # wangEditor v5 严格过滤不支持的标签(如 div)。为保证原样显示，不使用 div 包裹，直接返回 p 和 table 等基础标签
                wrapped_html = raw_html
                if not node.get("content"):
                    node["content"] = wrapped_html
                continue  # 成功提取 HTML，跳过纯文本提取
        
        # 降级或未匹配到 HTML，使用原有的纯文本提取逻辑
        start = find_title_pos(content_text, title, format_chapter_pos, format_chapter_end)
        
        if start < 0:
            if format_chapter_pos > 0:
                if not node.get("content"):
                    node["content"] = "【未精确匹配到该标题，以下为投标文件格式章节内容参考】\n\n" + content_text[format_chapter_pos:min(format_chapter_pos+2000, format_chapter_end)]
            else:
                if not node.get("content"):
                    node["content"] = ""
        else:
            end = format_chapter_end
            if next_title:
                found = find_title_pos(content_text, next_title, start + len(title), format_chapter_end)
                if found > start:
                    end = found
            extracted = content_text[start:end].strip("\n")
            if not node.get("content"):
                node["content"] = extracted

def _to_chinese_numeral(num: int) -> str:
    if num <= 0:
        return str(num)
    digits = "零一二三四五六七八九"
    if num < 10:
        return digits[num]
    if num < 20:
        return "十" + (digits[num % 10] if num % 10 else "")
    if num < 100:
        tens = num // 10
        ones = num % 10
        return digits[tens] + "十" + (digits[ones] if ones else "")
    if num < 1000:
        hundreds = num // 100
        remainder = num % 100
        if remainder == 0:
            return digits[hundreds] + "百"
        if remainder < 10:
            return digits[hundreds] + "百零" + digits[remainder]
        return digits[hundreds] + "百" + _to_chinese_numeral(remainder)
    thousands = num // 1000
    remainder = num % 1000
    if remainder == 0:
        return _to_chinese_numeral(thousands) + "千"
    if remainder < 100:
        return _to_chinese_numeral(thousands) + "千零" + _to_chinese_numeral(remainder)
    return _to_chinese_numeral(thousands) + "千" + _to_chinese_numeral(remainder)

def _strip_chapter_prefix(title: str) -> str:
    if not isinstance(title, str):
        return ""
    cleaned = title.strip()
    cleaned = re.sub(r"^\s*第[一二三四五六七八九十百千零〇两\d]+[章节部分篇卷][\s、.．:：-]*", "", cleaned)
    cleaned = re.sub(r"^\s*[（(]?\d+[）)](?:\.\d+)*[、.．:：-]?\s*", "", cleaned)
    cleaned = re.sub(r"^\s*\d+(?:\.\d+)*[、.．:：-]?\s*", "", cleaned)
    cleaned = re.sub(r"^\s*[一二三四五六七八九十百千零〇两]+[、.．:：-]\s*", "", cleaned)
    cleaned = re.sub(r"^\s*[（(][一二三四五六七八九十百千零〇两]+[）)][、.．:：-]?\s*", "", cleaned)
    return cleaned.strip()

def _build_chapter_prefix(path_indexes: List[int]) -> str:
    depth = len(path_indexes) - 1
    current = path_indexes[-1]
    if depth == 0:
        return f"{_to_chinese_numeral(current)}、"
    if depth == 1:
        return f"{current}."
    if depth == 2:
        return f"{path_indexes[1]}.{current}"
    return f"({current})"

def ensure_directory_numbering(directories) -> None:
    def walk(nodes, path_indexes: List[int]):
        if not isinstance(nodes, list):
            return
        for idx, node in enumerate(nodes, start=1):
            if not isinstance(node, dict):
                continue
            current_path = path_indexes + [idx]
            title = node.get("title", "")
            base_title = _strip_chapter_prefix(title) or str(title or "").strip()
            prefix = _build_chapter_prefix(current_path)
            node["title"] = f"{prefix} {base_title}".strip()
            children = node.get("children")
            if isinstance(children, list) and children:
                walk(children, current_path)
    walk(directories, [])

def replace_placeholders_in_directories(directories: list, elements_content: str) -> None:
    if not elements_content:
        return
    try:
        elements = json.loads(elements_content)
        replacements = []
        
        for category in elements:
            if category.get("title") == "关键信息与占位符映射":
                for subcat in category.get("subcategories", []):
                    for item in subcat.get("items", []):
                        placeholder = item.get("placeholder", "")
                        val = item.get("value") or item.get("description", "")
                        
                        if placeholder and val and len(val) >= 2:
                            # Filter out placeholder texts from prompt
                            if "提取到" not in val and "实际" not in val and "例如" not in val and "获取到" not in val:
                                replacements.append((val, placeholder))
        
        if not replacements:
            return

        # Sort replacements by length descending to replace longer strings first
        replacements.sort(key=lambda x: len(x[0]), reverse=True)

        def walk_and_replace(nodes):
            if not isinstance(nodes, list):
                return
            for node in nodes:
                if not isinstance(node, dict):
                    continue
                content = node.get("content", "")
                if content:
                    for val, placeholder in replacements:
                        content = content.replace(val, placeholder)
                    node["content"] = content
                
                children = node.get("children")
                if isinstance(children, list) and children:
                    walk_and_replace(children)

        walk_and_replace(directories)
    except Exception as e:
        logger.error(f"Failed to replace placeholders in directories: {e}")

@router.post("/{project_id}/generate-directories-stream")
async def generate_business_directories_stream(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    config = config_manager.load_config()
    if not config.get('api_key'):
        raise HTTPException(status_code=400, detail="请先配置OpenAI API密钥")

    openai_service = OpenAIService()

    async def generate():
        from app.database import SessionLocal
        import asyncio
        
        chapter_text = None
        if project.tender_content:
            _, _, chapter_text = get_format_chapter_info(project.tender_content)
            
        if chapter_text:
            system_prompt = """你是一名专业的标书编写专家。请根据提供的【投标文件格式】或【招标文件格式】章节内容，按其原本的目录层次生成一份标准的商务标目录大纲。
请严格输出JSON格式，并且【极其重要】：你必须在每个目录项的 `title` 字段前自动加上规范的章节编号（如：一、 / 1. / 1.1 / (1) 等），绝对不能只输出光秃秃的文本标题！
格式要求如下：
[
  {
    "id": "1",
    "title": "一、法定代表人身份证明",
    "description": "提供法定代表人身份证正反面复印件及身份证明书",
    "children": [
      {
        "id": "1.1",
        "title": "1. 身份证明书原件",
        "description": "..."
      }
    ]
  },
  ...
]
请确保直接输出JSON，不要包含任何其他说明文字、Markdown代码块标记。"""
            
            # 截取前8000字，避免超出token限制，通常目录结构在章节开头
            user_prompt = f"请根据以下提取的格式章节内容，提取其目录层次生成商务标目录大纲。\n\n【警告】：你输出的所有节点的 title 必须强制以正确的章节序号开头（即使原文中没有序号，你也要自动补齐！）。\n\n{chapter_text[:8000]}"
        else:
            system_prompt = """你是一名专业的标书编写专家。请根据提供的商务要求要素，生成一份标准的商务标目录大纲。
请严格输出JSON格式，并且【极其重要】：你必须在每个目录项的 `title` 字段前自动加上标准的章节编号（如：一、 / 1. / 1.1 / (1) 等），绝对不能只输出光秃秃的文本标题！
格式要求如下：
[
  {
    "id": "1",
    "title": "一、法定代表人身份证明",
    "description": "提供法定代表人身份证正反面复印件及身份证明书",
    "children": []
  },
  ...
]
目录结构通常包括：投标函、法定代表人身份证明、授权委托书、营业执照、资质证书、财务状况、业绩证明、人员证明等。
请确保直接输出JSON，不要包含任何其他说明文字、Markdown代码块标记。"""
            
            elements_str = project.elements_content or "{}"
            user_prompt = f"请根据以下商务要素，生成商务标目录大纲。\n\n【警告】：你输出的所有节点的 title 必须强制以正确的章节序号开头！\n\n{elements_str}"
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        db_gen = SessionLocal()
        try:
            full_response = ""
            try:
                async for chunk in openai_service.stream_chat_completion(messages, temperature=0.3):
                    full_response += chunk
                    yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
            except asyncio.CancelledError:
                logger.warning("Client disconnected during directory streaming. Saving partial result.")
            except Exception as inner_e:
                logger.error(f"Error during directory stream: {inner_e}")
                yield f"data: {json.dumps({'error': str(inner_e)}, ensure_ascii=False)}\n\n"
            
            project_gen = db_gen.query(BusinessBidProject).filter(BusinessBidProject.id == project_id).first()
            if project_gen:
                # 尝试清理可能存在的Markdown代码块标记
                cleaned_response = full_response.strip()
                if cleaned_response.startswith("```json"):
                    cleaned_response = cleaned_response[7:]
                if cleaned_response.endswith("```"):
                    cleaned_response = cleaned_response[:-3]
                cleaned_response = cleaned_response.strip()

                try:
                    generate_format_word_file(project_gen, db_gen)
                    logger.info("Successfully generated format word document.")
                except Exception as ex:
                    logger.error(f"Failed to generate format word document: {ex}")

                directories = None
                try:
                    directories = json.loads(cleaned_response)
                    ensure_directory_numbering(directories)
                    cleaned_response = json.dumps(directories, ensure_ascii=False)
                except Exception as ex:
                    logger.error(f"Failed to normalize generated directories: {ex}")

                if directories is not None and project_gen.tender_content:
                    try:
                        extract_content_for_directories(directories, project_gen)
                        replace_placeholders_in_directories(directories, project_gen.elements_content)
                        cleaned_response = json.dumps(directories, ensure_ascii=False)
                    except Exception as ex:
                        logger.error(f"Failed to pre-extract content for directories: {ex}")

                # Save the result to DB
                project_gen.directories_content = cleaned_response
                project_gen.status = "generating"
                db_gen.commit()
                
                logger.info("Successfully saved directories to DB.")
            
            yield "data: [DONE]\n\n"
        except asyncio.CancelledError:
            logger.warning("Client disconnected. Directory generation aborted.")
        except Exception as e:
            logger.error(f"Directory generation failed: {e}")
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"
            yield "data: [DONE]\n\n"
        finally:
            db_gen.close()

    return sse_response(generate())

@router.get("/{project_id}/directories")
def get_business_directories(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    if project.directories_content:
        try:
            directories = json.loads(project.directories_content)
            original_dump = json.dumps(directories, ensure_ascii=False, sort_keys=True)
            ensure_directory_numbering(directories)
            normalized_dump = json.dumps(directories, ensure_ascii=False, sort_keys=True)
            if normalized_dump != original_dump:
                project.directories_content = json.dumps(directories, ensure_ascii=False)
                db.commit()
            return {"directories": directories}
        except json.JSONDecodeError:
            pass
            
    return {
        "directories": []
    }

@router.get("/{project_id}/directory-source")
def get_directory_source_content(project_id: str, node_title: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if not project.tender_content:
        raise HTTPException(status_code=400, detail="Tender document not uploaded or content is empty")

    title = (node_title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="node_title is required")

    directories = []
    if project.directories_content:
        try:
            directories = json.loads(project.directories_content) or []
        except json.JSONDecodeError:
            directories = []

    flat: List[tuple[str, int]] = []

    def flatten(nodes, depth: int = 0):
        if not isinstance(nodes, list):
            return
        for n in nodes:
            if not isinstance(n, dict):
                continue
            t = str(n.get("title") or "")
            flat.append((t, depth))
            children = n.get("children") or []
            flatten(children, depth + 1)

    flatten(directories, 0)

    idx: Optional[int] = None
    current_depth = 0
    for i, (t, d) in enumerate(flat):
        if t == title:
            idx = i
            current_depth = d
            break

    next_title: Optional[str] = None
    if idx is not None:
        for j in range(idx + 1, len(flat)):
            t, d = flat[j]
            if d <= current_depth and t:
                next_title = t
                break
        if next_title is None and idx + 1 < len(flat):
            t, _ = flat[idx + 1]
            if t:
                next_title = t

    content_text = project.tender_content

    # 1. 定位“投标文件格式”大章节
    format_chapter_pos, format_chapter_end, _ = get_format_chapter_info(content_text)
    if format_chapter_pos < 0:
        format_chapter_pos = 0
        format_chapter_end = len(content_text)

    # 只在“投标文件格式”章节内查找对应的目录节点
    start = find_title_pos(content_text, title, format_chapter_pos, format_chapter_end)

    # 尝试读取新生成的格式文件 URL
    document_url = project.tender_document_url
    if project.other_urls:
        try:
            other_urls = json.loads(project.other_urls)
            if "format_document_url" in other_urls:
                document_url = other_urls["format_document_url"]
        except Exception:
            pass

    if start < 0:
        # 如果在格式章节内没找到，直接返回格式章节的开头部分供用户参考
        if format_chapter_pos > 0:
            page_num = 1
            matches = list(re.finditer(r'--- 第 (\d+) 页 ---', content_text[:format_chapter_pos]))
            if matches:
                page_num = int(matches[-1].group(1))
            return {
                "success": True, 
                "matched": False, 
                "node_title": title, 
                "content": "【未精确匹配到该标题，以下为投标文件格式章节内容参考】\n\n" + content_text[format_chapter_pos:min(format_chapter_pos+2000, format_chapter_end)],
                "page_num": page_num,
                "document_url": document_url
            }
        return {"success": True, "matched": False, "node_title": title, "content": ""}

    # 尝试寻找页码 (向后查找最近的页码标记)
    page_num = 1
    page_marker_pattern = re.compile(r'--- 第 (\d+) 页 ---')
    matches = list(page_marker_pattern.finditer(content_text[:start]))
    if matches:
        page_num = int(matches[-1].group(1))

    end = format_chapter_end
    if next_title:
        found = find_title_pos(content_text, next_title, start + len(title), format_chapter_end)
        if found > start:
            end = found

    extracted = content_text[start:end].strip("\n")
    
    # 尝试读取新生成的格式文件 URL
    document_url = project.tender_document_url
    if project.other_urls:
        try:
            other_urls = json.loads(project.other_urls)
            if "format_document_url" in other_urls:
                document_url = other_urls["format_document_url"]
        except Exception:
            pass

    return {
        "success": True, 
        "matched": True, 
        "node_title": title, 
        "next_title": next_title, 
        "content": extracted,
        "page_num": page_num,
        "document_url": document_url
    }

class SmartFillRequest(BaseModel):
    html_content: str
    resources: list[dict]
    node_title: str | None = None
    node_description: str | None = None
    tender_requirement: str | None = None

def _strip_html_text(html_content: str) -> str:
    if not html_content:
        return ""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")
        text = soup.get_text("\n")
        return re.sub(r"\n{3,}", "\n\n", text).strip()
    except Exception:
        return re.sub(r"<[^>]+>", " ", html_content or "").strip()

def _normalize_title_for_match(text: str) -> str:
    t = (text or "").strip()
    t = re.sub(r"^\s*第[一二三四五六七八九十百千零〇两\d]+[章节部分篇卷][\s、.．:：-]*", "", t)
    t = re.sub(r"^\s*[（(]?\d+[）)](?:\.\d+)*[、.．:：-]?\s*", "", t)
    t = re.sub(r"^\s*\d+(?:\.\d+)*[、.．:：-]?\s*", "", t)
    t = re.sub(r"^\s*[一二三四五六七八九十百千零〇两]+[、.．:：-]\s*", "", t)
    t = re.sub(r"^\s*[（(][一二三四五六七八九十百千零〇两]+[）)][、.．:：-]?\s*", "", t)
    return t.strip()

def _plain_text_to_html(text: str) -> str:
    lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    html_parts: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            html_parts.append("<p><br/></p>")
        else:
            html_parts.append(f"<p>{escape(stripped)}</p>")
    return "".join(html_parts) if html_parts else "<p><br/></p>"

def _generated_text_to_html(text: str) -> str:
    content = (text or "").strip()
    if not content:
        return "<p><br/></p>"
    if re.search(r"<[a-zA-Z][^>]*>", content):
        return content
    lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    html_parts: list[str] = []
    i = 0
    sep_pattern = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
    while i < len(lines):
        line = lines[i].strip()
        if line and "|" in line and i + 1 < len(lines) and sep_pattern.match(lines[i + 1].strip()):
            header_cells = [c.strip() for c in line.strip("|").split("|")]
            body_rows: list[list[str]] = []
            i += 2
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line or "|" not in row_line:
                    break
                body_rows.append([c.strip() for c in row_line.strip("|").split("|")])
                i += 1
            table_html = "<table><thead><tr>" + "".join([f"<th>{escape(c)}</th>" for c in header_cells]) + "</tr></thead><tbody>"
            for row in body_rows:
                table_html += "<tr>" + "".join([f"<td>{escape(c)}</td>" for c in row]) + "</tr>"
            table_html += "</tbody></table>"
            html_parts.append(table_html)
            continue
        if not line:
            html_parts.append("<p><br/></p>")
        else:
            html_parts.append(f"<p>{escape(line)}</p>")
        i += 1
    return "".join(html_parts) if html_parts else "<p><br/></p>"

def _merge_generated_after_title(original_html: str, generated_html: str, node_title: str | None) -> str:
    html_content = original_html or "<p><br/></p>"
    if not generated_html.strip():
        return html_content
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")
        generated_soup = BeautifulSoup(generated_html, "html.parser")
        old_block = soup.find("div", attrs={"data-ai-custom-fill": "1"})
        if old_block:
            old_block.decompose()
        wrapper = soup.new_tag("div")
        wrapper["data-ai-custom-fill"] = "1"
        for child in list(generated_soup.contents):
            wrapper.append(child)
        target_title = _normalize_title_for_match(node_title or "")
        title_tags = ["h1", "h2", "h3", "h4", "h5", "h6", "p", "strong", "span", "div"]
        title_node = None
        if target_title:
            for tag in soup.find_all(title_tags):
                text = _normalize_title_for_match(tag.get_text(" ", strip=True))
                if text and (text == target_title or target_title in text):
                    title_node = tag
                    break
        if title_node is not None:
            title_node.insert_after(wrapper)
        elif soup.body:
            soup.body.insert(0, wrapper)
        else:
            soup.append(wrapper)
        if soup.body:
            return soup.body.decode_contents()
        return str(soup)
    except Exception:
        return f"{html_content}{generated_html}"

@router.post("/{project_id}/smart-fill")
async def smart_fill_template(project_id: str, request: SmartFillRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    openai_service = OpenAIService()
    
    # 构建资源文本
    resources_text = "\n\n".join([f"【{r.get('title', '')}】\n{r.get('content', '')}" for r in request.resources])
    
    # 构建项目关键信息
    project_info = ""
    if project.elements_content:
        try:
            elements = json.loads(project.elements_content)
            for category in elements:
                if "占位符" in category.get("title", "") or "关键信息" in category.get("title", ""):
                    project_info = json.dumps(category, ensure_ascii=False)
                    break
        except Exception:
            pass

    html_content = request.html_content or ""
    plain_content = _strip_html_text(html_content)
    is_custom_format = bool(re.search(r"格式\s*自拟", plain_content))
    tender_requirement = (request.tender_requirement or "").strip()
    if is_custom_format and not tender_requirement and request.node_title:
        try:
            source_info = get_directory_source_content(project_id, request.node_title, current_user, db)
            if isinstance(source_info, dict):
                tender_requirement = (source_info.get("content") or "").strip()
        except Exception:
            tender_requirement = ""

    if is_custom_format:
        system_prompt = """你是一个专业的技术标书编写专家。请基于章节信息和招标要求，生成可直接用于投标文件正文的内容。
要求：
1. 严格围绕章节标题、章节描述和招标文件要求生成，不写无关内容。
2. 语言正式、专业、可执行，不空泛，不口语化。
3. 不得编造未提供的事实数据；缺失信息可用“按招标文件要求执行”这类合规表述。
4. 只输出章节正文内容，不要输出标题。
5. 输出必须是HTML片段，段落用<p>，如需表格必须使用<table>/<tr>/<td>真实HTML标签，禁止输出Markdown表格语法（如 |---|）。"""
        user_prompt = f"""
【章节标题】：
{request.node_title or ""}

【章节描述】：
{request.node_description or ""}

【招标文件对本章节的要求】：
{tender_requirement}

【可用于写作的资料】：
{resources_text}

【当前模板内容】：
{plain_content}

请直接输出本章节正文：
"""
    else:
        system_prompt = """你是一个智能的文档填充助手。你的任务是将提供的相关资料和项目关键信息，精准地填入给定的 HTML 模板中。
核心要求：
1. 保持原有 HTML 结构和格式绝对不变（特别是表格 <table>、加粗 <strong>、下划线 <u>、以及我预设的外层包裹 div 等）。
2. 识别 HTML 中的下划线填空区（如 `____`）、括号（如 `（  ）`）、或者我设定的占位符（如 `{{PROJECT_NAME}}`）。
3. 根据提供的资料内容，将相关信息填充到这些空白处，替换掉下划线或占位符。如果原本有下划线，填充后可以保留适量下划线或者直接填入文字，但不能破坏整体美观和结构。
4. 如果提供的资料中没有可以对应填充的信息，则必须保留原有空白、下划线或占位符，绝对不要瞎编。
5. 只返回处理后的完整 HTML 内容，不要包含 ```html 等任何 Markdown 标记。"""
        user_prompt = f"""
【项目关键信息与占位符映射】：
{project_info}

【可用于填充的资料】：
{resources_text}

【待填充的 HTML 模板】：
{html_content}

请输出填充后的 HTML 模板内容：
"""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    try:
        logger.info(f"Starting smart fill for project {project_id}")
        response_text = ""
        async for chunk in openai_service.stream_chat_completion(messages, temperature=0.1):
            response_text += chunk
            
        cleaned_html = response_text.strip()
        if is_custom_format:
            generated_text = cleaned_html
            if generated_text.startswith("```"):
                generated_text = re.sub(r"^```[a-zA-Z]*\s*", "", generated_text)
                generated_text = re.sub(r"\s*```$", "", generated_text)
            generated_html = _generated_text_to_html(generated_text)
            cleaned_html = _merge_generated_after_title(html_content, generated_html, request.node_title)
        else:
            if cleaned_html.startswith("```html"):
                cleaned_html = cleaned_html[7:]
            if cleaned_html.endswith("```"):
                cleaned_html = cleaned_html[:-3]
            
        logger.info(f"Smart fill completed for project {project_id}")
        return {"success": True, "filled_content": cleaned_html.strip()}
    except Exception as e:
        logger.error(f"Smart fill failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class BusinessProjectUpdate(BaseModel):
    status: str | None = None
    tender_document_name: str | None = None
    tender_content: str | None = None
    elements_content: str | None = None
    directories_content: str | None = None

@router.put("/{project_id}")
def update_business_project(project_id: str, data: BusinessProjectUpdate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    if data.status is not None:
        project.status = data.status
    if data.tender_document_name is not None:
        project.tender_document_name = data.tender_document_name
    if data.tender_content is not None:
        project.tender_content = data.tender_content
    if data.elements_content is not None:
        project.elements_content = data.elements_content
    if data.directories_content is not None:
        try:
            directories = json.loads(data.directories_content)
            ensure_directory_numbering(directories)
            project.directories_content = json.dumps(directories, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Failed to process directories_content on PUT update: {e}")
            project.directories_content = data.directories_content
        
    db.commit()
    return {"success": True, "message": "Project updated"}

@router.post("/{project_id}/mark-completed")
def mark_business_project_completed(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    project.status = "completed"
    db.commit()
    
    return {"success": True, "message": "Project marked as completed", "status": project.status}

@router.get("/{project_id}/match-resource")
def match_business_resource(project_id: str, node_title: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    results = []
    
    # 简单的基于关键词的匹配逻辑
    title = node_title.lower()

    def s(v):
        return v or ""

    if any(k in title for k in ["授权委托", "被授权", "委托人身份证", "授权人身份证", "授权身份证", "受托人身份证"]):
        info = db.query(CompanyInfo).filter(CompanyInfo.user_id == current_user.id).first()
        if info:
            results.append({
                "id": f"auth_idcard_{info.id}",
                "type": "company",
                "title": "授权委托人身份证",
                "content": "\n".join([
                    f"姓名: {s(info.authorized_person)}",
                    f"性别: {s(info.authorized_person_gender)}",
                    f"出生日期: {s(info.authorized_person_birth_date)}",
                    f"身份证号: {s(info.authorized_person_id_number)}",
                    f"有效期自: {s(info.authorized_person_id_valid_from)}",
                    f"有效期至: {s(info.authorized_person_id_valid_to)}",
                    f"长期有效: {'是' if info.authorized_person_id_long_term else '否'}" if info.authorized_person_id_long_term is not None else "",
                ]).strip(),
                "image_url": info.authorized_person_id_card_url,
            })

    if any(k in title for k in ["法人身份证", "法定代表人身份证", "法人代表身份证"]) or ("法人" in title and "身份证" in title):
        info = db.query(CompanyInfo).filter(CompanyInfo.user_id == current_user.id).first()
        if info:
            results.append({
                "id": f"legal_idcard_{info.id}",
                "type": "company",
                "title": "法定代表人身份证",
                "content": "\n".join([
                    f"姓名: {s(info.legal_person)}",
                    f"性别: {s(info.legal_person_gender)}",
                    f"出生日期: {s(info.legal_person_birth_date)}",
                    f"身份证号: {s(info.legal_person_id_number)}",
                    f"有效期自: {s(info.legal_person_id_valid_from)}",
                    f"有效期至: {s(info.legal_person_id_valid_to)}",
                    f"长期有效: {'是' if info.legal_person_id_long_term else '否'}" if info.legal_person_id_long_term is not None else "",
                ]).strip(),
                "image_url": info.legal_person_id_card_url,
            })
    
    if any(k in title for k in ["营业执照", "法人", "公司", "主体"]):
        info = db.query(CompanyInfo).filter(CompanyInfo.user_id == current_user.id).first()
        if info:
            results.append({
                "id": f"company_{info.id}",
                "type": "company",
                "title": info.company_name or "公司信息",
                "content": f"公司名称: {s(info.company_name)}\n法定代表人: {s(info.legal_person)}\n注册资本: {s(info.registered_capital)}\n成立日期: {s(info.establish_date)}",
                "image_url": info.business_license_url
            })
            
    if any(k in title for k in ["资质", "许可", "证书"]):
        quals = db.query(Qualification).filter(Qualification.user_id == current_user.id).all()
        for q in quals:
            results.append({
                "id": f"qual_{q.id}",
                "type": "qualification",
                "title": q.cert_name,
                "content": f"证书名称: {q.cert_name}\n证书编号: {q.cert_number}\n发证机关: {q.issue_org}\n有效期至: {q.valid_end_date}",
                "image_url": q.cert_image_url
            })
            
    if any(k in title for k in ["人员", "经理", "负责人", "团队"]):
        personnel = db.query(Personnel).filter(Personnel.user_id == current_user.id).all()
        for p in personnel:
            results.append({
                "id": f"person_{p.id}",
                "type": "personnel",
                "title": f"{p.name} - {p.position or p.title or '人员'}",
                "content": f"姓名: {p.name}\n职务: {p.position}\n职称: {p.title}\n学历: {p.education}",
                "image_url": p.cert_image_url or p.id_card_url
            })
            
    if any(k in title for k in ["财务", "审计", "报表", "资信"]):
        fins = db.query(FinancialInfo).filter(FinancialInfo.user_id == current_user.id).all()
        for f in fins:
            results.append({
                "id": f"fin_{f.id}",
                "type": "financial",
                "title": f.info_name,
                "content": f"财务信息名称: {f.info_name}\n类型: {f.info_type}\n金额: {f.amount}",
                "image_url": f.file_url
            })
            
    if any(k in title for k in ["业绩", "合同", "项目", "经验"]):
        perfs = db.query(Performance).filter(Performance.user_id == current_user.id).all()
        for p in perfs:
            results.append({
                "id": f"perf_{p.id}",
                "type": "performance",
                "title": p.project_name,
                "content": f"项目名称: {p.project_name}\n客户名称: {p.client_name}\n合同金额: {p.contract_amount}\n项目时间: {p.start_date} 至 {p.end_date}",
                "image_url": p.contract_url or p.acceptance_url
            })
            
    # 如果没有匹配到任何数据，返回空列表
    return {"success": True, "results": results}


@router.post("/{project_id}/verify-stream")
async def verify_business_bid_stream(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    config = config_manager.load_config()
    if not config.get('api_key'):
        raise HTTPException(status_code=400, detail="请先配置OpenAI API密钥")

    openai_service = OpenAIService()

    # 1. 组装待核验的数据
    tender_text = project.tender_content or ""
    directories_text = project.directories_content or ""
    
    # 获取公司信息
    company_info = db.query(CompanyInfo).filter(CompanyInfo.user_id == current_user.id).first()
    company_name = company_info.company_name if company_info else "未配置公司名称"

    system_prompt = f"""你是一名资深的商务标审核专家。你的任务是对用户填写的商务标内容进行合规性和准确性核验。
【核验原则】：
1. 合规性优先：提取招标文件中的所有废标条款（如“未加盖公章废标”、“漏填保证金废标”），并在填充内容中寻找违规风险。
2. 对“签字盖章”、“份数要求”、“装订要求”等硬性规定生成醒目的“待盖章项清单”提示。
3. 信息精准性：项目名称、招标编号、投标截止时间等核心信息必须100%匹配。
4. 检查公司名称拼写是否准确（本公司名称：{company_name}）。
5. 检查金额单位是否统一（例如招标要求“万元”，实际填写“元”需指出）。

请流式输出核验结果，每次发现一个问题或提示，就输出一个 JSON 对象。
必须严格按以下 JSON 格式输出，不要包含 ```json 等任何 Markdown 标记：
{{"type": "danger|warning|info", "category": "废标风险|盖章签字|信息校验|硬性要求", "message": "具体说明", "matched_text": "原文中的问题片段", "node_id": "对应目录节点的id（如果能推断出来）"}}

注意：只输出合法的 JSON 对象，每个对象占一行，或者确保它们是独立的结构。为了流式解析方便，请在每个 JSON 对象后加一个换行符 `\\n`。"""

    user_prompt = f"""
【招标文件片段（用于提取要求）】：
{tender_text[:30000]}...

【已填写的商务标内容】：
{directories_text[:50000]}...

请开始核验：
"""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    async def generate():
        try:
            buffer = ""
            async for chunk in openai_service.stream_chat_completion(messages, temperature=0.1):
                buffer += chunk
                # 尝试逐行解析 JSON
                lines = buffer.split('\n')
                if len(lines) > 1:
                    for line in lines[:-1]:
                        line = line.strip()
                        if line:
                            try:
                                # 去除可能包含的逗号
                                if line.endswith(','):
                                    line = line[:-1]
                                parsed = json.loads(line)
                                yield f"data: {json.dumps({'result': parsed}, ensure_ascii=False)}\n\n"
                            except json.JSONDecodeError:
                                pass
                    buffer = lines[-1]
            
            # 处理最后一行
            if buffer.strip():
                try:
                    line = buffer.strip()
                    if line.endswith(','):
                        line = line[:-1]
                    parsed = json.loads(line)
                    yield f"data: {json.dumps({'result': parsed}, ensure_ascii=False)}\n\n"
                except json.JSONDecodeError:
                    pass
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error(f"Error during AI verification: {e}")
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

    return sse_response(generate())


from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import io
import re
import base64
from urllib.request import urlopen
from fastapi.responses import StreamingResponse
from urllib.parse import quote
from ..models.schemas import WordExportRequest

def set_run_font_simsun(run) -> None:
    run.font.name = "宋体"
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), "宋体")

@router.post("/{project_id}/export-docx")
def export_business_docx(project_id: str, request: WordExportRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(BusinessBidProject).filter(BusinessBidProject.id == project_id, BusinessBidProject.user_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    company_info = db.query(CompanyInfo).filter(CompanyInfo.user_id == current_user.id).first()
    company_name = company_info.company_name if company_info else "投标人全称"

    try:
        doc = docx.Document()
        
        # 1. 页面设置 (A4, 边距上2.5下2.5左2.0右2.0)
        section = doc.sections[0]
        section.page_height = docx.shared.Mm(297)
        section.page_width = docx.shared.Mm(210)
        section.top_margin = docx.shared.Cm(2.5)
        section.bottom_margin = docx.shared.Cm(2.5)
        section.left_margin = docx.shared.Cm(2.0)
        section.right_margin = docx.shared.Cm(2.0)

        # 2. 字体与行距基础设置
        styles = doc.styles
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = "宋体"
        normal_font.size = Pt(12)  # 小四号
        if normal_style._element.rPr is None:
            normal_style._element._add_rPr()
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        
        normal_format = normal_style.paragraph_format
        normal_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.EXACTLY
        normal_format.line_spacing = Pt(22)  # 固定值20-22pt
        
        # 3. 封面
        for _ in range(5):
            doc.add_paragraph()
        
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(f"{project.project_name} 商务投标书")
        title_run.font.name = "黑体"
        title_run.font.size = Pt(26) # 一号字
        title_run.bold = True
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        
        for _ in range(15):
            doc.add_paragraph()
            
        company_p = doc.add_paragraph()
        company_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_p.add_run(f"投标人：{company_name}")
        company_run.font.name = "黑体"
        company_run.font.size = Pt(15) # 小三号字
        company_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        import datetime
        date_str = datetime.datetime.now().strftime("%Y年%m月%d日")
        date_run = date_p.add_run(f"日期：{date_str}")
        date_run.font.name = "宋体"
        date_run.font.size = Pt(15) # 小三号字
        date_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        
        doc.add_page_break()

        # 4. 页眉页脚设置
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_p.add_run(f"{project.project_name} 商务投标书 {company_name}")
        header_run.font.name = "宋体"
        header_run.font.size = Pt(9) # 小五号
        header_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run("投标书正文")
        footer_run.font.name = "宋体"
        footer_run.font.size = Pt(9)
        footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 5. Markdown 解析逻辑
        def add_markdown_runs(para, text: str):
            pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`|!\[.*?\]\(.*?\))"
            parts = re.split(pattern, text)
            for part in parts:
                if not part: continue
                run = para.add_run()
                if part.startswith("![") and "](" in part and part.endswith(")"):
                    try:
                        m = re.match(r"!\[(.*?)\]\((.*?)\)", part)
                        if m:
                            alt_text, img_url = m.groups()
                            img_stream = None
                            if img_url.startswith("data:image"):
                                header, encoded = img_url.split(",", 1)
                                img_stream = io.BytesIO(base64.b64decode(encoded))
                            elif img_url.startswith("http"):
                                response = urlopen(img_url, timeout=5)
                                img_stream = io.BytesIO(response.read())
                            if img_stream:
                                run.add_picture(img_stream, width=Inches(6))
                                continue
                    except Exception: pass
                if part.startswith("**") and part.endswith("**") and len(part) > 4:
                    run.text = part[2:-2]
                    run.bold = True
                elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run.text = part[1:-1]
                    run.italic = True
                elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                    run.text = part[1:-1]
                else:
                    run.text = part
                set_run_font_simsun(run)

        def parse_markdown_blocks(content: str):
            blocks = []
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].rstrip("\r").strip()
                if not line:
                    i += 1
                    continue
                if line.startswith("- ") or line.startswith("* ") or re.match(r"^\d+\.\s", line):
                    items = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        if stripped.startswith("- ") or stripped.startswith("* "):
                            text = re.sub(r"^[-*]\s+", "", stripped).strip()
                            if text: items.append(("unordered", None, text))
                            i += 1
                            continue
                        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                        if m_num:
                            num_str, text = m_num.groups()
                            if text.strip(): items.append(("ordered", num_str, text.strip()))
                            i += 1
                            continue
                        break
                    if items: blocks.append(("list", items))
                    continue
                if "|" in line:
                    rows = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        if "|" in stripped:
                            if not re.match(r"^\|?[-\s\|]+\|?$", stripped):
                                cells = [c.strip() for c in stripped.split("|")]
                                if cells and not cells[0]: cells.pop(0)
                                if cells and not cells[-1]: cells.pop()
                                if cells: rows.append(cells)
                            i += 1
                        else: break
                    if rows: blocks.append(("table", rows))
                    continue
                if line.startswith("#"):
                    m = re.match(r"^(#+)\s*(.*)$", line)
                    if m:
                        level_marks, title_text = m.groups()
                        blocks.append(("heading", min(len(level_marks), 3), title_text.strip()))
                    i += 1
                    continue
                para_lines = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if stripped and not stripped.startswith("-") and not stripped.startswith("*") and "|" not in stripped and not stripped.startswith("#"):
                        para_lines.append(stripped)
                        i += 1
                    else: break
                if para_lines:
                    blocks.append(("paragraph", " ".join(para_lines)))
                else: i += 1
            return blocks

        def render_markdown_blocks(blocks):
            for block in blocks:
                kind = block[0]
                if kind == "list":
                    for item_kind, num_str, text in block[1]:
                        p = doc.add_paragraph()
                        run = p.add_run("• " if item_kind == "unordered" else f"{num_str}. ")
                        set_run_font_simsun(run)
                        add_markdown_runs(p, text)
                elif kind == "table":
                    rows = block[1]
                    if rows:
                        max_cols = max(len(r) for r in rows)
                        if max_cols > 0:
                            table = doc.add_table(rows=len(rows), cols=max_cols)
                            table.style = 'Table Grid'
                            for i, row in enumerate(rows):
                                for j, cell_text in enumerate(row):
                                    if j < max_cols:
                                        cell = table.cell(i, j)
                                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                                        add_markdown_runs(p, cell_text)
                            doc.add_paragraph()
                elif kind == "heading":
                    _, level, text = block
                    heading = doc.add_heading(text, level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for hr in heading.runs: set_run_font_simsun(hr)
                elif kind == "paragraph":
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = docx.shared.Cm(0.74)
                    add_markdown_runs(p, block[1])

        # 正文写入
        def write_nodes(items, depth=1):
            if not isinstance(items, list):
                return
            for item in items:
                title = item.title if hasattr(item, 'title') else item.get('title', '')
                content = item.content if hasattr(item, 'content') else item.get('content', '')
                
                # 写入标题
                p = doc.add_paragraph()
                run = p.add_run(title)
                if depth == 1:
                    run.font.name = "黑体"
                    run.font.size = Pt(22)
                    run.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                elif depth == 2:
                    run.font.name = "黑体"
                    run.font.size = Pt(16)
                    run.bold = True
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                elif depth == 3:
                    run.font.name = "楷体"
                    run.font.size = Pt(14)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
                else:
                    run.font.name = "宋体"
                    run.font.size = Pt(12)
                    run.bold = True
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    
                if content:
                    blocks = parse_markdown_blocks(content)
                    render_markdown_blocks(blocks)
                
                children = item.children if hasattr(item, 'children') else item.get('children', [])
                if children:
                    write_nodes(children, depth + 1)

        write_nodes(request.outline)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{project.project_name or '商务标文件'}.docx"
        encoded_filename = quote(filename)
        content_disposition = f"attachment; filename*=UTF-8''{encoded_filename}"
        headers = {"Content-Disposition": content_disposition}

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        logger.error(f"Failed to export docx: {e}")
        raise HTTPException(status_code=500, detail=str(e))
